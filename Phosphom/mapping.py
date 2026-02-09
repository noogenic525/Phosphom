"""
Step 1 — Phosphorylation site mapping in Word documents.

Reads a Word document containing full protein sequences and peptide
fragments with phosphorylation probability markers (e.g. ``S(0.477)``),
identifies matching kinases via the motif database, and produces a
colour-highlighted annotated document.
"""

from __future__ import annotations

import logging
import re
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from .motif_db import identify_kinases

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
AMINO_ACIDS = frozenset("ACDEFGHIKLMNPQRSTVWY")
MIN_SEQUENCE_LENGTH: int = 30  # residues — handles small proteins (e.g. ubiquitin)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def clean_all_whitespace(text: str) -> str:
    """Remove every whitespace character from *text*."""
    return re.sub(r"\s+", "", text)


def is_protein_sequence(text: str) -> bool:
    """Decide whether *text* looks like a full protein amino-acid sequence.

    Criteria
    --------
    1. After whitespace removal, length ≥ ``MIN_SEQUENCE_LENGTH``.
    2. ≥ 90 % of characters are standard amino acids.
    3. No parenthesised numbers (peptide-annotation artefact).
    """
    cleaned = clean_all_whitespace(text).upper()
    if len(cleaned) < MIN_SEQUENCE_LENGTH:
        return False
    aa_count = sum(1 for c in cleaned if c in AMINO_ACIDS)
    if aa_count / len(cleaned) < 0.90:
        return False
    # Parenthesised digits indicate an annotated peptide, not a raw sequence
    if re.search(r"\(\d", text):
        return False
    return True


def apply_colors_to_full_seq(
    para, text: str, color_map: List[Optional[int]]
) -> None:
    """Apply per-residue colour highlighting to a full-sequence paragraph.

    Groups consecutive residues sharing the same colour into single Word
    runs for compact document rendering.
    """
    if not para or not color_map:
        return
    para.clear()
    current_run_text = ""
    current_color = color_map[0]
    for i in range(len(text)):
        if color_map[i] == current_color:
            current_run_text += text[i]
        else:
            run = para.add_run(current_run_text)
            run.font.highlight_color = current_color
            current_run_text = text[i]
            current_color = color_map[i]
    if current_run_text:
        run = para.add_run(current_run_text)
        run.font.highlight_color = current_color


# ---------------------------------------------------------------------------
# Main mapping function
# ---------------------------------------------------------------------------
def start_mapping(
    input_path: str,
    output_path: str,
    min_confidence: float = 0.0,
) -> int:
    """Map phosphorylation sites to kinases in a Word document.

    Iterates through paragraphs, detects full protein sequences and their
    peptide fragments, performs motif matching for each phospho-site, then
    writes an annotated + colour-highlighted output document.

    Parameters
    ----------
    input_path : str
        Path to the input Word (.docx) file.
    output_path : str
        Path for the output annotated Word (.docx) file.
    min_confidence : float, optional
        Minimum specificity score for kinase matches (default 0.0).

    Returns
    -------
    int
        Number of peptide lines processed.
    """
    doc = Document(input_path)
    current_full_seq_text: str = ""
    current_full_seq_para = None
    current_color_map: List[Optional[int]] = []
    processed_count: int = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # --- Detect full protein sequence lines ---
        if is_protein_sequence(text):
            # Flush the previous sequence's colouring
            if current_full_seq_para:
                apply_colors_to_full_seq(
                    current_full_seq_para,
                    current_full_seq_text,
                    current_color_map,
                )
            current_full_seq_text = clean_all_whitespace(text).upper()
            current_full_seq_para = para
            current_color_map = [None] * len(current_full_seq_text)
            logger.debug(
                "Detected protein sequence: %d residues", len(current_full_seq_text)
            )
            continue

        # --- Process peptide lines with phospho-site markers ---
        if "(" in text and current_full_seq_text:
            # Remove trailing position annotation like (123, 456)
            original_peptide = re.sub(r"\s*\([\d,\s]+\)$", "", text).strip()
            pure_peptide = clean_all_whitespace(
                re.sub(r"\(.*?\)", "", original_peptide)
            ).upper()
            start_idx = current_full_seq_text.find(pure_peptide)

            if start_idx == -1:
                logger.warning(
                    "Peptide not found in current sequence: %s", pure_peptide[:30]
                )
                continue

            # Colour the peptide region gray (non-phospho residues)
            for i in range(start_idx, start_idx + len(pure_peptide)):
                if i < len(current_color_map) and (
                    current_color_map[i] != WD_COLOR_INDEX.YELLOW
                ):
                    current_color_map[i] = WD_COLOR_INDEX.GRAY_25

            # Split peptide by probability markers: e.g. "ABCS(0.477)EFK"
            results_labels: List[str] = []
            temp_idx = 0
            parts = re.split(r"(\([0-9.]+\))", original_peptide)

            for part in parts:
                if re.match(r"\([0-9.]+\)", part):
                    real_idx = start_idx + temp_idx
                    phospho_pos = real_idx - 1

                    # Bounds check — prevents silent bug when marker is at pos 0
                    if 0 <= phospho_pos < len(current_color_map):
                        current_color_map[phospho_pos] = WD_COLOR_INDEX.YELLOW

                        kinase_hits = identify_kinases(
                            current_full_seq_text, phospho_pos, min_confidence
                        )
                        label = str(real_idx)
                        if kinase_hits:
                            kinase_strs = [
                                f"{name}:{score:.2f}"
                                for name, score in kinase_hits
                            ]
                            label += " " + ", ".join(kinase_strs)
                        results_labels.append(label)
                    else:
                        logger.warning(
                            "Phospho-site position out of bounds: %d "
                            "(sequence length: %d)",
                            phospho_pos,
                            len(current_color_map),
                        )
                else:
                    temp_idx += len(clean_all_whitespace(part))

            # Rewrite paragraph with colour highlighting
            para.clear()
            for i, part in enumerate(parts):
                if re.match(r"\([0-9.]+\)", part):
                    run = para.add_run(part)
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                else:
                    if i + 1 < len(parts) and re.match(
                        r"\([0-9.]+\)", parts[i + 1]
                    ):
                        if part:  # guard against empty string
                            para.add_run(part[:-1])
                            run = para.add_run(part[-1])
                            run.font.highlight_color = WD_COLOR_INDEX.RED
                    else:
                        para.add_run(part)

            para.add_run(f" ({', '.join(results_labels)})")
            processed_count += 1

    # Flush the last sequence's colouring
    if current_full_seq_para:
        apply_colors_to_full_seq(
            current_full_seq_para,
            current_full_seq_text,
            current_color_map,
        )

    doc.save(output_path)
    logger.info("Mapping complete: %d peptide lines processed", processed_count)
    return processed_count
