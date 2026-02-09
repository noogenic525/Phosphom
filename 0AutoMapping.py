import re
import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def clean_all_whitespace(text):
    return re.sub(r'\s+', '', text)

def apply_colors_to_full_seq(para, text, color_map):
    if not para or not color_map: return
    para.clear()
    if not text: return
    current_run_text = ""
    current_color = color_map[0]
    for i in range(len(text)):
        if color_map[i] == current_color:
            current_run_text += text[i]
        else:
            run = para.add_run(current_run_text)
            run.font.highlight_color = current_color
            current_run_text = text[i]
            current_color = color_map[i]
    run = para.add_run(current_run_text)
    run.font.highlight_color = current_color

def start_mapping(input_path, output_path):
    doc = Document(input_path)
    current_full_seq_text = ""
    current_full_seq_para = None
    current_color_map = []
    processed_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        if len(text) > 50 and text.isupper() and "(" not in text:
            if current_full_seq_para:
                apply_colors_to_full_seq(current_full_seq_para, current_full_seq_text, current_color_map)
            current_full_seq_text = clean_all_whitespace(text)
            current_full_seq_para = para
            current_color_map = [None] * len(current_full_seq_text)
            continue
        if "(" in text and current_full_seq_text:
            original_peptide = re.sub(r'\s*\([\d,\s]+\)$', '', text).strip()
            pure_peptide = clean_all_whitespace(re.sub(r'\(.*?\)', '', original_peptide))
            start_idx = current_full_seq_text.find(pure_peptide)
            if start_idx != -1:
                for i in range(start_idx, start_idx + len(pure_peptide)):
                    if current_color_map[i] != WD_COLOR_INDEX.YELLOW:
                        current_color_map[i] = WD_COLOR_INDEX.GRAY_25
                p_indices = []
                temp_idx = 0
                parts = re.split(r'(\([0-9.]+\))', original_peptide)
                for part in parts:
                    if re.match(r'\([0-9.]+\)', part):
                        current_color_map[start_idx + temp_idx - 1] = WD_COLOR_INDEX.YELLOW
                        p_indices.append(start_idx + temp_idx)
                    else:
                        temp_idx += len(clean_all_whitespace(part))
                para.clear()
                for i, part in enumerate(parts):
                    if re.match(r'\([0-9.]+\)', part):
                        para.add_run(part).font.highlight_color = WD_COLOR_INDEX.YELLOW
                    else:
                        if i + 1 < len(parts) and re.match(r'\([0-9.]+\)', parts[i+1]):
                            para.add_run(part[:-1])
                            para.add_run(part[-1]).font.highlight_color = WD_COLOR_INDEX.RED
                        else:
                            para.add_run(part)
                para.add_run(f" ({', '.join(map(str, p_indices))})")
                processed_count += 1
    if current_full_seq_para:
        apply_colors_to_full_seq(current_full_seq_para, current_full_seq_text, current_color_map)
    doc.save(output_path)
    return processed_count

# --- GUI 화면 구성 ---
def run_app():
    root = tk.Tk()
    root.title("인산화 서열 매핑 자동화 도구")
    root.geometry("400x200")

    def select_file():
        file_path = filedialog.askopenfilename(title="작업할 Word 파일을 선택하세요", filetypes=[("Word files", "*.docx")])
        if not file_path: return
        
        save_path = filedialog.asksaveasfilename(title="결과를 저장할 이름을 입력하세요", defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if not save_path: return

        try:
            count = start_mapping(file_path, save_path)
            messagebox.showinfo("성공", f"작업이 완료되었습니다!\n매핑된 항목: {count}개")
        except Exception as e:
            messagebox.showerror("오류", f"작업 중 오류가 발생했습니다:\n{str(e)}")

    tk.Label(root, text="인산화 서열 매핑 프로그램", font=("Arial", 16, "bold"), pady=20).pack()
    tk.Button(root, text="파일 선택 및 실행하기", command=select_file, bg="blue", fg="white", width=20, height=2).pack()
    root.mainloop()

if __name__ == "__main__":
    run_app()