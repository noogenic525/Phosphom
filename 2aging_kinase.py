import re
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# --- [1] 노화 연구 특화 모티프 데이터베이스 (Regex 기반) ---
MOTIF_DB = [
    # 1. Proline-directed Kinases (중첩 발생 빈도 높음)
    ("mTOR&CDK", r"[ST]P"),
    ("Erk", r"P.[ST]P"),
    ("p38&MAPK", r"[LIVM].[ST]P"),
    ("GSK-3beta&CDK", r"[ST].{3}[ST]P"),
    ("GSK-3beta", r"[ST].{2}[ST]"),

    # 2. Basophilic Kinases (R/K 염기성 아미노산 공유)
    ("PKA", r"[RK][RK].[ST]"),
    ("CaMK2", r"R..[ST]"),
    ("AMPK", r"[LIVM].[RK].[ST]"),

    # 3. Acidic Kinases
    ("CK2", r"[ST]..[DE]"),

    # 4. 매우 특이적인 패턴
    ("IKK", r"DSG..[ST]"),
    ("ATM&ATR", r"[ST]Q"),
]

def identify_kinase(full_seq, index):
    """
    인산화 위치(index)를 중심으로 모든 모티프를 검사하여 중첩된 Kinase를 모두 반환합니다.
    """
    start = max(0, index - 5)
    end = min(len(full_seq), index + 6)
    fragment = full_seq[start:end]
    
    matches = []
    
    # 모든 패턴을 순회하며 매칭되는 모든 이름을 리스트에 추가
    for name, pattern in MOTIF_DB:
        if re.search(pattern, fragment):
            matches.append(name)
    
    # 매칭된 것이 있다면 쉼표로 구분된 문자열 반환, 없으면 None
    return ", ".join(matches) if matches else None

# --- [2] 데이터 처리 및 매핑 로직 ---

def clean_all_whitespace(text):
    return re.sub(r'\s+', '', text)

def apply_colors_to_full_seq(para, text, color_map):
    if not para or not color_map: return
    para.clear()
    current_run_text = ""
    current_color = color_map[0]
    for i in range(len(text)):
        if color_map[i] == current_color:
            current_run_text += text[i]
        else:
            run = para.add_run(current_run_text)
            run.font.highlight_color = current_color
            current_run_text = text[i]
            current_color = color_map[i]
    run = para.add_run(current_run_text)
    run.font.highlight_color = current_color

def start_mapping(input_path, output_path):
    doc = Document(input_path)
    current_full_seq_text = ""
    current_full_seq_para = None
    current_color_map = []
    processed_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        
        if len(text) > 50 and text.isupper() and "(" not in text:
            if current_full_seq_para:
                apply_colors_to_full_seq(current_full_seq_para, current_full_seq_text, current_color_map)
            current_full_seq_text = clean_all_whitespace(text)
            current_full_seq_para = para
            current_color_map = [None] * len(current_full_seq_text)
            continue
            
        if "(" in text and current_full_seq_text:
            original_peptide = re.sub(r'\s*\([\d,\s]+\)$', '', text).strip()
            pure_peptide = clean_all_whitespace(re.sub(r'\(.*?\)', '', original_peptide))
            start_idx = current_full_seq_text.find(pure_peptide)
            
            if start_idx != -1:
                for i in range(start_idx, start_idx + len(pure_peptide)):
                    if current_color_map[i] != WD_COLOR_INDEX.YELLOW:
                        current_color_map[i] = WD_COLOR_INDEX.GRAY_25
                
                results_labels = [] 
                temp_idx = 0
                parts = re.split(r'(\([0-9.]+\))', original_peptide)
                
                for part in parts:
                    if re.match(r'\([0-9.]+\)', part):
                        real_idx = start_idx + temp_idx
                        current_color_map[real_idx - 1] = WD_COLOR_INDEX.YELLOW
                        
                        # --- [중첩 분석 수행] ---
                        kinase = identify_kinase(current_full_seq_text, real_idx - 1)
                        label = str(real_idx)
                        if kinase:
                            label += f" {kinase}"
                        results_labels.append(label)
                    else:
                        temp_idx += len(clean_all_whitespace(part))
                
                para.clear()
                for i, part in enumerate(parts):
                    if re.match(r'\([0-9.]+\)', part):
                        para.add_run(part).font.highlight_color = WD_COLOR_INDEX.YELLOW
                    else:
                        if i + 1 < len(parts) and re.match(r'\([0-9.]+\)', parts[i+1]):
                            para.add_run(part[:-1])
                            para.add_run(part[-1]).font.highlight_color = WD_COLOR_INDEX.RED
                        else:
                            para.add_run(part)
                
                para.add_run(f" ({', '.join(results_labels)})")
                processed_count += 1
                
    if current_full_seq_para:
        apply_colors_to_full_seq(current_full_seq_para, current_full_seq_text, current_color_map)
    doc.save(output_path)
    return processed_count

def run_app():
    root = tk.Tk()
    root.title("인산화 모티프 분석 도구 (중첩 허용)")
    root.geometry("400x200")

    def select_file():
        file_path = filedialog.askopenfilename(title="Word 파일 선택", filetypes=[("Word files", "*.docx")])
        if not file_path: return
        save_path = filedialog.asksaveasfilename(title="결과 저장", defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if not save_path: return
        try:
            count = start_mapping(file_path, save_path)
            messagebox.showinfo("성공", f"작업 완료! 처리된 항목: {count}개\n중첩 매칭이 적용되었습니다.")
        except Exception as e:
            messagebox.showerror("오류", f"에러 발생:\n{str(e)}")

    tk.Label(root, text="Kinase Motif Mapping", font=("Arial", 16, "bold"), pady=20).pack()
    tk.Button(root, text="파일 선택 및 실행", command=select_file, bg="blue", fg="white", width=20, height=2).pack()
    root.mainloop()

if __name__ == "__main__":
    run_app()